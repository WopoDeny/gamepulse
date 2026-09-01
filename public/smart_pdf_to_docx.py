"""
smart_pdf_to_docx.py

Production-oriented hybrid PDF -> editable DOCX converter.

The module intentionally keeps the high-quality OCR route from the original
working version:

    PaddleOCR-VL -> result.save_to_word(...)

The main performance improvements are architectural rather than replacing the
OCR model:

1. Fully digital PDFs use one whole-document pdf2docx conversion. OCR is never
   imported or loaded for those files.
2. PaddleOCR-VL is lazy-loaded only when a real scan/image-only page is found.
3. The OCR pipeline is cached in RAM and reused for every following conversion
   in the same Python process.
4. A persistent JSON-lines worker is included for integration from non-Python
   applications. Start it once, keep it alive, and avoid reloading the 0.9B
   model for every PDF.
5. Third-party console noise is suppressed as much as possible; application
   logs report only useful stages and timings.

Recommended environment:
    Windows 10/11, 64-bit Python 3.9-3.13

CPU installation:
    python -m pip install paddlepaddle==3.2.1 -i https://www.paddlepaddle.org.cn/packages/stable/cpu/
    python -m pip install -U "paddleocr[doc-parser]" pdf2docx pymupdf python-docx docxcompose

Simple CLI:
    python smart_pdf_to_docx.py input.pdf output.docx

Download/cache OCR models once:
    python smart_pdf_to_docx.py --prepare-ocr

Persistent worker:
    python smart_pdf_to_docx.py --worker --preload-ocr
"""

from __future__ import annotations

# ---------------------------------------------------------------------------
# Standard library imports and quiet third-party defaults.
# These environment variables are set before Paddle is imported.
# ---------------------------------------------------------------------------

import argparse
import contextlib
import io
import json
import logging
import os
import shutil
import sys
import tempfile
import threading
import time
import warnings
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Callable, Iterable, Optional, Sequence

os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
os.environ.setdefault("FLAGS_minloglevel", "3")
os.environ.setdefault("GLOG_minloglevel", "3")
os.environ.setdefault("PADDLE_LOG_LEVEL", "ERROR")
os.environ.setdefault("DNNL_VERBOSE", "0")
os.environ.setdefault("ONEDNN_VERBOSE", "0")
os.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "1")
os.environ.setdefault("TRANSFORMERS_VERBOSITY", "error")
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

warnings.filterwarnings("ignore", message=r".*fitz.*deprecated.*")
warnings.filterwarnings("ignore", message=r".*Non compatible API.*")
warnings.filterwarnings("ignore", message=r".*To copy construct from a tensor.*")

# Use the modern PyMuPDF import but keep the familiar "fitz" alias internally.
import pymupdf as fitz
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Mm


ProgressCallback = Callable[[int, int, str], None]

LOGGER = logging.getLogger("smart_pdf_to_docx")
LOGGER.addHandler(logging.NullHandler())


# ---------------------------------------------------------------------------
# Public configuration and result data.
# ---------------------------------------------------------------------------

@dataclass(slots=True)
class ConverterConfig:
    """Runtime options for PDF -> DOCX conversion.

    The defaults favor the original high-quality OCR behavior while keeping
    normal digital PDFs on a much faster route.
    """

    # Routing mode:
    #   auto    -> inspect each page and choose digital/OCR automatically.
    #   digital -> never use OCR.
    #   ocr     -> force PaddleOCR-VL for every page.
    mode: str = "auto"

    # A page with a real native text layer is normally treated as digital.
    minimum_native_text_chars: int = 20
    minimum_native_words: int = 4

    # If there is not enough native text, these values help distinguish a scan.
    full_page_image_threshold: float = 0.62

    # Optional compatibility switch. Disabled by default because an ordinary
    # digital page containing a photo/logo should not pay the OCR cost.
    ocr_image_heavy_digital_pages: bool = False
    image_area_ocr_threshold: float = 0.40

    # PaddleOCR-VL settings. These are intentionally close to the first version
    # because that route produced the best DOCX output in testing.
    device: Optional[str] = None
    engine: Optional[str] = None
    use_doc_orientation_classify: bool = True
    use_doc_unwarping: bool = True
    use_layout_detection: bool = True
    use_ocr_for_image_block: bool = True
    use_seal_recognition: bool = True
    use_chart_recognition: bool = False
    layout_shape_mode: str = "auto"

    # If a layout-enabled OCR result contains almost no text, retry only those
    # weak pages without layout detection and keep whichever result is better.
    retry_without_layout: bool = True
    retry_if_recognized_chars_below: int = 24

    # File/output behavior.
    overwrite: bool = True
    fallback_digital_to_ocr: bool = True
    validate_output: bool = True
    preserve_styles_when_merging: bool = True

    # Paddle CPU threads. None lets Paddle choose its own default.
    cpu_threads: Optional[int] = None


@dataclass(slots=True)
class PageAnalysis:
    """Routing information for one PDF page."""

    page_index: int
    width_mm: float
    height_mm: float
    text_chars: int
    word_count: int
    image_area_ratio: float
    largest_image_ratio: float
    route: str
    reason: str


@dataclass(slots=True)
class BatchResult:
    """Result of one file inside convert_many()."""

    input_pdf: str
    output_docx: Optional[str]
    success: bool
    elapsed_seconds: float
    error: Optional[str] = None


# ---------------------------------------------------------------------------
# Global OCR cache.
#
# This cache is the key to fast integration. Once PaddleOCR-VL is loaded in a
# running process, every later scan reuses the same pipeline object.
# ---------------------------------------------------------------------------

_OCR_PIPELINE_CACHE: dict[tuple, object] = {}
_OCR_PIPELINE_LOCK = threading.RLock()


# ---------------------------------------------------------------------------
# Logging helpers.
# ---------------------------------------------------------------------------

def configure_logging(level: str = "INFO") -> None:
    """Configure concise console logs for standalone use.

    Library users may skip this function and configure Python logging in their
    own application instead.
    """

    numeric_level = getattr(logging, level.upper(), logging.INFO)
    LOGGER.setLevel(numeric_level)

    # Avoid duplicate handlers if configure_logging() is called more than once.
    for handler in list(LOGGER.handlers):
        if getattr(handler, "_smart_pdf_handler", False):
            LOGGER.removeHandler(handler)

    handler = logging.StreamHandler(sys.stderr)
    handler._smart_pdf_handler = True  # type: ignore[attr-defined]
    handler.setLevel(numeric_level)
    handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s", "%H:%M:%S"))
    LOGGER.addHandler(handler)
    LOGGER.propagate = False


@contextlib.contextmanager
def _quiet_library_output(enabled: bool = True):
    """Hide noisy Python-level stdout/stderr produced by Paddle/PaddleX.

    Native C/C++ libraries may still emit a rare line directly to the process
    file descriptor, but the large model-creation/configuration dumps are
    normally suppressed by this context plus the environment variables above.
    """

    if not enabled:
        yield
        return

    stdout_buffer = io.StringIO()
    stderr_buffer = io.StringIO()
    with contextlib.redirect_stdout(stdout_buffer), contextlib.redirect_stderr(stderr_buffer):
        yield


# ---------------------------------------------------------------------------
# PDF inspection and routing.
# ---------------------------------------------------------------------------

def _emit(callback: Optional[ProgressCallback], current: int, total: int, message: str) -> None:
    if callback is not None:
        callback(current, total, message)


def _is_pdf(path: Path) -> bool:
    if path.suffix.lower() != ".pdf":
        return False
    try:
        with path.open("rb") as file:
            return file.read(5) == b"%PDF-"
    except OSError:
        return False


def _rect_area(rect: fitz.Rect) -> float:
    return max(0.0, float(rect.width)) * max(0.0, float(rect.height))


def _page_image_metrics(page: fitz.Page) -> tuple[float, float]:
    """Return (total image area ratio, largest image area ratio)."""

    page_area = max(_rect_area(page.rect), 1.0)
    try:
        blocks = page.get_text("dict").get("blocks", [])
    except Exception:
        blocks = []

    total_image_area = 0.0
    largest_image_area = 0.0

    for block in blocks:
        if block.get("type") != 1:
            continue
        bbox = block.get("bbox")
        if not bbox or len(bbox) != 4:
            continue
        rect = fitz.Rect(bbox) & page.rect
        area = _rect_area(rect)
        if area <= 0:
            continue
        total_image_area += area
        largest_image_area = max(largest_image_area, area)

    return (
        min(total_image_area / page_area, 1.0),
        min(largest_image_area / page_area, 1.0),
    )


def _native_text_metrics(page: fitz.Page) -> tuple[int, int]:
    """Return native text character count and native word count."""

    try:
        text = page.get_text("text") or ""
        words = page.get_text("words") or []
    except Exception:
        return 0, 0

    normalized = "".join(ch for ch in text if not ch.isspace())
    return len(normalized), len(words)


def _choose_page_route(
    text_chars: int,
    word_count: int,
    image_area_ratio: float,
    largest_image_ratio: float,
    config: ConverterConfig,
) -> tuple[str, str]:
    """Choose the conversion route for one page.

    Important optimization:
    if the page already has a healthy native text layer, it stays on the
    digital route even if it contains photos/logos. This keeps ordinary PDFs
    with text and tables fast.
    """

    mode = config.mode.lower().strip()

    if mode == "digital":
        return "digital", "forced digital mode"
    if mode == "ocr":
        return "ocr", "forced OCR mode"
    if mode != "auto":
        raise ValueError("config.mode must be 'auto', 'digital', or 'ocr'.")

    has_native_text = (
        text_chars >= config.minimum_native_text_chars
        and word_count >= config.minimum_native_words
    )

    # Sparse native pages (cover sheets, short forms, table-only pages) are also
    # cheap digital candidates when they are not dominated by one scan image.
    has_sparse_native_text = (
        text_chars >= 8
        and word_count >= 2
        and largest_image_ratio < config.full_page_image_threshold
    )

    if has_native_text or has_sparse_native_text:
        if (
            config.ocr_image_heavy_digital_pages
            and image_area_ratio >= config.image_area_ocr_threshold
        ):
            return "ocr", "image-heavy digital page; OCR explicitly enabled"
        return "digital", "usable native PDF text layer"

    if largest_image_ratio >= config.full_page_image_threshold:
        return "ocr", "scan-like full-page image"

    return "ocr", "insufficient native PDF text"


def analyze_pdf(
    input_pdf: str | os.PathLike,
    config: Optional[ConverterConfig] = None,
) -> list[PageAnalysis]:
    """Inspect every page without loading OCR models."""

    config = config or ConverterConfig()
    input_path = Path(input_pdf).expanduser().resolve()

    if not input_path.exists():
        raise FileNotFoundError(input_path)
    if not _is_pdf(input_path):
        raise ValueError(f"Not a valid PDF file: {input_path}")

    analyses: list[PageAnalysis] = []
    pdf = fitz.open(str(input_path))

    try:
        for page_index, page in enumerate(pdf):
            text_chars, word_count = _native_text_metrics(page)
            image_area_ratio, largest_image_ratio = _page_image_metrics(page)
            route, reason = _choose_page_route(
                text_chars,
                word_count,
                image_area_ratio,
                largest_image_ratio,
                config,
            )

            analyses.append(
                PageAnalysis(
                    page_index=page_index,
                    width_mm=float(page.rect.width) * 25.4 / 72.0,
                    height_mm=float(page.rect.height) * 25.4 / 72.0,
                    text_chars=text_chars,
                    word_count=word_count,
                    image_area_ratio=image_area_ratio,
                    largest_image_ratio=largest_image_ratio,
                    route=route,
                    reason=reason,
                )
            )
    finally:
        pdf.close()

    if not analyses:
        raise ValueError("The PDF contains no pages.")

    return analyses


# ---------------------------------------------------------------------------
# Fast digital conversion.
# ---------------------------------------------------------------------------

def _pdf2docx_converter_class():
    """Import pdf2docx lazily so OCR-only processes do not need it immediately."""

    from pdf2docx import Converter
    return Converter


def _convert_digital_pdf(input_pdf: Path, output_docx: Path) -> None:
    """Convert an entire native/digital PDF in one pdf2docx call."""

    Converter = _pdf2docx_converter_class()
    converter = Converter(str(input_pdf))
    try:
        converter.convert(str(output_docx))
    finally:
        converter.close()

    if not output_docx.exists() or output_docx.stat().st_size == 0:
        raise RuntimeError("pdf2docx did not create a valid output file.")


def _extract_single_page_pdf(source_pdf: fitz.Document, page_index: int, output_pdf: Path) -> None:
    """Create a temporary one-page PDF for mixed digital/OCR documents."""

    single = fitz.open()
    try:
        single.insert_pdf(source_pdf, from_page=page_index, to_page=page_index)
        single.save(str(output_pdf), garbage=4, deflate=True)
    finally:
        single.close()


def _convert_digital_page(single_page_pdf: Path, output_docx: Path) -> None:
    _convert_digital_pdf(single_page_pdf, output_docx)


# ---------------------------------------------------------------------------
# PaddleOCR-VL loading and persistent in-process cache.
# ---------------------------------------------------------------------------

def _paddle_dependency_error(error: Exception) -> RuntimeError:
    return RuntimeError(
        "OCR is required, but PaddleOCR-VL could not be loaded.\n\n"
        "Install the OCR stack first:\n"
        "python -m pip install paddlepaddle==3.2.1 "
        "-i https://www.paddlepaddle.org.cn/packages/stable/cpu/\n"
        'python -m pip install -U "paddleocr[doc-parser]" '
        "pdf2docx pymupdf python-docx docxcompose\n\n"
        f"Original error: {error}"
    )


def _ocr_pipeline_key(config: ConverterConfig) -> tuple:
    return (
        config.device,
        config.engine,
        config.use_doc_orientation_classify,
        config.use_doc_unwarping,
        config.use_layout_detection,
        config.use_ocr_for_image_block,
        config.use_seal_recognition,
        config.use_chart_recognition,
        config.cpu_threads,
    )


def _get_ocr_pipeline(config: ConverterConfig, *, quiet_library_logs: bool = True):
    """Return a cached PaddleOCR-VL pipeline, loading it only once per process."""

    key = _ocr_pipeline_key(config)

    # Fast path without locking after the model is already initialized.
    cached = _OCR_PIPELINE_CACHE.get(key)
    if cached is not None:
        return cached

    # The lock prevents two application threads from loading two copies of the
    # 0.9B model at the same time.
    with _OCR_PIPELINE_LOCK:
        cached = _OCR_PIPELINE_CACHE.get(key)
        if cached is not None:
            return cached

        started = time.perf_counter()
        LOGGER.info("OCR model: loading PaddleOCR-VL (first use in this process)...")

        try:
            with _quiet_library_output(quiet_library_logs):
                from paddleocr import PaddleOCRVL

                kwargs = {
                    "use_doc_orientation_classify": config.use_doc_orientation_classify,
                    "use_doc_unwarping": config.use_doc_unwarping,
                    "use_layout_detection": config.use_layout_detection,
                    "use_ocr_for_image_block": config.use_ocr_for_image_block,
                    "use_seal_recognition": config.use_seal_recognition,
                    "use_chart_recognition": config.use_chart_recognition,
                    "use_queues": True,
                }

                if config.device:
                    kwargs["device"] = config.device
                if config.engine:
                    kwargs["engine"] = config.engine
                if config.cpu_threads is not None:
                    kwargs["cpu_threads"] = int(config.cpu_threads)

                pipeline = PaddleOCRVL(**kwargs)

        except Exception as error:
            raise _paddle_dependency_error(error) from error

        _OCR_PIPELINE_CACHE[key] = pipeline
        LOGGER.info("OCR model: ready in %.1f s", time.perf_counter() - started)
        return pipeline


def prepare_ocr(
    config: Optional[ConverterConfig] = None,
    *,
    quiet_library_logs: bool = True,
) -> None:
    """Download/load OCR models now instead of waiting for the first scan.

    One-time installation use:
        python smart_pdf_to_docx.py --prepare-ocr

    Production application use:
        call prepare_ocr() once when the long-running application/worker starts.

    Important: model FILES stay cached on disk between runs. The 0.9B model
    still has to be loaded from disk into RAM once after each process start.
    """

    config = config or ConverterConfig()
    _get_ocr_pipeline(config, quiet_library_logs=quiet_library_logs)


def clear_ocr_model_cache() -> None:
    """Drop Python references to loaded OCR pipelines.

    Normally do not call this in a bulk conversion application, because keeping
    the model in RAM is exactly what makes later scan conversions faster.
    """

    with _OCR_PIPELINE_LOCK:
        _OCR_PIPELINE_CACHE.clear()


# ---------------------------------------------------------------------------
# OCR result helpers.
# ---------------------------------------------------------------------------

def _result_json(result) -> dict:
    """Return a stable dictionary across PaddleOCR/PaddleX result wrappers."""

    try:
        value = result.json
    except Exception:
        return {}

    if not isinstance(value, dict):
        return {}

    wrapped = value.get("res")
    return wrapped if isinstance(wrapped, dict) else value


def _result_markdown_text(result) -> str:
    try:
        markdown = result.markdown
        if not isinstance(markdown, dict):
            return ""
        value = markdown.get("markdown_texts", "")
        if isinstance(value, str):
            return value
        if isinstance(value, (list, tuple)):
            return "\n".join(str(item) for item in value)
        return str(value or "")
    except Exception:
        return ""


def _recursive_text_size(value) -> int:
    if isinstance(value, str):
        return len("".join(ch for ch in value if not ch.isspace()))

    if isinstance(value, dict):
        total = 0
        for key, child in value.items():
            if key in {"block_content", "text", "rec_text", "rec_texts", "markdown_texts", "content"}:
                total += _recursive_text_size(child)
            elif isinstance(child, (dict, list, tuple)):
                total += _recursive_text_size(child)
        return total

    if isinstance(value, (list, tuple)):
        return sum(_recursive_text_size(item) for item in value)

    return 0


def _recognized_char_score(result) -> int:
    markdown_score = len(
        "".join(ch for ch in _result_markdown_text(result) if not ch.isspace())
    )
    json_score = _recursive_text_size(_result_json(result))
    return max(markdown_score, json_score)


def _result_input_path(result) -> Optional[Path]:
    data = _result_json(result)
    raw = data.get("input_path")
    if not raw:
        return None
    try:
        return Path(str(raw)).resolve()
    except Exception:
        return Path(str(raw))


def _map_results_to_inputs(inputs: Sequence[Path], results: Iterable) -> dict[Path, object]:
    normalized_inputs = [path.resolve() for path in inputs]
    output: dict[Path, object] = {}
    unassigned_results = []
    by_name: dict[str, list[Path]] = {}

    for path in normalized_inputs:
        by_name.setdefault(path.name.lower(), []).append(path)

    for result in results:
        result_path = _result_input_path(result)
        matched = None

        if result_path is not None:
            candidate = result_path.resolve()
            if candidate in normalized_inputs:
                matched = candidate
            else:
                candidates = by_name.get(candidate.name.lower(), [])
                if len(candidates) == 1:
                    matched = candidates[0]

        if matched is not None and matched not in output:
            output[matched] = result
        else:
            unassigned_results.append(result)

    remaining_inputs = [path for path in normalized_inputs if path not in output]
    for path, result in zip(remaining_inputs, unassigned_results):
        output[path] = result

    return output


def _predict_ocr(
    pipeline,
    inputs: Sequence[Path],
    config: ConverterConfig,
    *,
    use_layout_detection: bool,
    quiet_library_logs: bool = True,
    progress_callback: Optional[ProgressCallback] = None,
) -> list:
    """Run PaddleOCR-VL while keeping its internal queue/batch behavior.

    We iterate the returned generator instead of immediately wrapping it in
    list(...), so a concise progress update can be emitted whenever Paddle
    yields a page result.
    """

    if not inputs:
        return []

    kwargs = {
        "input": [str(path) for path in inputs],
        "use_doc_orientation_classify": config.use_doc_orientation_classify,
        "use_doc_unwarping": config.use_doc_unwarping,
        "use_layout_detection": use_layout_detection,
        "use_chart_recognition": config.use_chart_recognition,
        "use_seal_recognition": config.use_seal_recognition,
        "use_ocr_for_image_block": config.use_ocr_for_image_block,
        "layout_shape_mode": config.layout_shape_mode,
        "use_queues": True,
    }

    results = []
    total = len(inputs)
    started = time.perf_counter()

    with _quiet_library_output(quiet_library_logs):
        iterator = pipeline.predict(**kwargs)
        for result in iterator:
            results.append(result)
            done = len(results)
            _emit(progress_callback, done, total, f"OCR result {done}/{total} ready")

    LOGGER.info("OCR inference: %d page(s) in %.1f s", total, time.perf_counter() - started)
    return results


def _best_ocr_results(
    inputs: Sequence[Path],
    config: ConverterConfig,
    *,
    quiet_library_logs: bool = True,
    progress_callback: Optional[ProgressCallback] = None,
) -> dict[Path, object]:
    """Run the original high-quality OCR route and optional weak-result retry."""

    pipeline = _get_ocr_pipeline(config, quiet_library_logs=quiet_library_logs)

    primary_results = _predict_ocr(
        pipeline,
        inputs,
        config,
        use_layout_detection=config.use_layout_detection,
        quiet_library_logs=quiet_library_logs,
        progress_callback=progress_callback,
    )
    primary_map = _map_results_to_inputs(inputs, primary_results)

    missing = [path.resolve() for path in inputs if path.resolve() not in primary_map]
    if missing:
        raise RuntimeError(
            "PaddleOCR-VL did not return results for: "
            + ", ".join(path.name for path in missing)
        )

    if not config.retry_without_layout or not config.use_layout_detection:
        return primary_map

    retry_inputs = [
        path.resolve()
        for path in inputs
        if _recognized_char_score(primary_map[path.resolve()])
        < config.retry_if_recognized_chars_below
    ]

    if not retry_inputs:
        return primary_map

    LOGGER.info("OCR retry: %d weak page(s) without layout detection", len(retry_inputs))

    retry_results = _predict_ocr(
        pipeline,
        retry_inputs,
        config,
        use_layout_detection=False,
        quiet_library_logs=quiet_library_logs,
        progress_callback=None,
    )
    retry_map = _map_results_to_inputs(retry_inputs, retry_results)

    for path in retry_inputs:
        retry_result = retry_map.get(path)
        if retry_result is None:
            continue
        if _recognized_char_score(retry_result) > _recognized_char_score(primary_map[path]):
            primary_map[path] = retry_result

    return primary_map


# ---------------------------------------------------------------------------
# OCR -> DOCX export. This deliberately preserves the first version's
# result.save_to_word(...) behavior because it produced the best layout.
# ---------------------------------------------------------------------------

def _save_ocr_result_to_docx(result, output_docx: Path, scratch_dir: Path) -> None:
    save_dir = scratch_dir / f"word_{output_docx.stem}"
    save_dir.mkdir(parents=True, exist_ok=True)
    before = set(save_dir.glob("*.docx"))

    result.save_to_word(save_path=str(save_dir))

    after = set(save_dir.glob("*.docx"))
    created = sorted(
        after - before,
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )

    if not created:
        created = sorted(after, key=lambda path: path.stat().st_mtime, reverse=True)

    if not created:
        raise RuntimeError(
            "PaddleOCR-VL recognized the page but did not export a DOCX file."
        )

    shutil.copy2(created[0], output_docx)


def _force_page_dimensions(docx_path: Path, width_mm: float, height_mm: float) -> None:
    document = Document(str(docx_path))
    for section in document.sections:
        section.page_width = Mm(width_mm)
        section.page_height = Mm(height_mm)
    document.save(str(docx_path))


def _append_page_break(docx_path: Path) -> None:
    document = Document(str(docx_path))
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.save(str(docx_path))


def _merge_docx_pages(
    page_docx_files: Sequence[Path],
    output_docx: Path,
    config: ConverterConfig,
) -> None:
    if not page_docx_files:
        raise RuntimeError("There are no page DOCX files to merge.")

    if len(page_docx_files) == 1:
        shutil.copy2(page_docx_files[0], output_docx)
        return

    try:
        from docxcompose.composer import Composer
    except Exception as error:
        raise RuntimeError(
            "docxcompose is required to combine converted pages. "
            "Install it with: python -m pip install docxcompose"
        ) from error

    for page_docx in page_docx_files[:-1]:
        _append_page_break(page_docx)

    master = Document(str(page_docx_files[0]))
    try:
        composer = Composer(master, preserve_styles=config.preserve_styles_when_merging)
    except TypeError:
        composer = Composer(master)

    for page_docx in page_docx_files[1:]:
        composer.append(Document(str(page_docx)))

    composer.save(str(output_docx))


def _validate_docx(path: Path) -> None:
    if not path.exists():
        raise RuntimeError(f"DOCX output does not exist: {path}")
    if path.stat().st_size < 1000:
        raise RuntimeError(f"DOCX output is unexpectedly small: {path.stat().st_size} bytes")
    if not zipfile.is_zipfile(path):
        raise RuntimeError("The produced file is not a valid DOCX ZIP package.")
    Document(str(path))


# ---------------------------------------------------------------------------
# Main public conversion API.
# ---------------------------------------------------------------------------

def convert_pdf_to_docx(
    input_pdf: str | os.PathLike,
    output_docx: str | os.PathLike | None = None,
    *,
    config: Optional[ConverterConfig] = None,
    progress_callback: Optional[ProgressCallback] = None,
    quiet_library_logs: bool = True,
) -> Path:
    """Convert one PDF to editable DOCX.

    AUTO mode:
      - healthy native PDF pages -> pdf2docx;
      - scan/image-only pages -> original PaddleOCR-VL + save_to_word route;
      - failed native page conversion -> OCR fallback;
      - OCR pipeline remains cached in RAM for future calls in this process.
    """

    started_at = time.perf_counter()
    config = config or ConverterConfig()
    input_path = Path(input_pdf).expanduser().resolve()

    if not input_path.exists():
        raise FileNotFoundError(input_path)
    if not _is_pdf(input_path):
        raise ValueError(f"Not a valid PDF file: {input_path}")

    output_path = (
        input_path.with_suffix(".docx")
        if output_docx is None
        else Path(output_docx).expanduser().resolve()
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if output_path.exists():
        if config.overwrite:
            output_path.unlink()
        else:
            raise FileExistsError(output_path)

    analyses = analyze_pdf(input_path, config=config)
    total_pages = len(analyses)
    digital_count = sum(1 for item in analyses if item.route == "digital")
    ocr_count = total_pages - digital_count

    LOGGER.info(
        "PDF: %s | pages=%d | digital=%d | OCR=%d",
        input_path.name,
        total_pages,
        digital_count,
        ocr_count,
    )
    _emit(progress_callback, 0, total_pages, f"analyzed: digital={digital_count}, ocr={ocr_count}")

    # ------------------------------------------------------------------
    # Fastest and most common path: the entire document already contains
    # usable native text. No Paddle import, no model loading, no page split.
    # ------------------------------------------------------------------
    if ocr_count == 0:
        LOGGER.info("Fast path: whole-document digital conversion (OCR not loaded)")
        _convert_digital_pdf(input_path, output_path)

        if config.validate_output:
            _validate_docx(output_path)

        elapsed = time.perf_counter() - started_at
        LOGGER.info("Done: %s (%.1f s)", output_path, elapsed)
        _emit(progress_callback, total_pages, total_pages, f"done in {elapsed:.1f}s")
        return output_path

    # ------------------------------------------------------------------
    # Hybrid/scanned document. Only now do we split pages and load OCR.
    # ------------------------------------------------------------------
    with tempfile.TemporaryDirectory(prefix="smart_pdf_to_docx_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        pages_dir = temp_dir / "pages"
        word_dir = temp_dir / "word"
        ocr_export_dir = temp_dir / "ocr_exports"

        pages_dir.mkdir(parents=True, exist_ok=True)
        word_dir.mkdir(parents=True, exist_ok=True)
        ocr_export_dir.mkdir(parents=True, exist_ok=True)

        source_pdf = fitz.open(str(input_path))
        try:
            single_page_pdfs: dict[int, Path] = {}
            for analysis in analyses:
                single_pdf = pages_dir / f"page_{analysis.page_index + 1:05d}.pdf"
                _extract_single_page_pdf(source_pdf, analysis.page_index, single_pdf)
                single_page_pdfs[analysis.page_index] = single_pdf
        finally:
            source_pdf.close()

        page_docx_paths: dict[int, Path] = {}
        ocr_page_indexes: list[int] = []

        # Convert digital pages first. If a native conversion unexpectedly
        # fails, that page is moved to the OCR list rather than failing the PDF.
        for analysis in analyses:
            page_number = analysis.page_index + 1

            if analysis.route == "ocr":
                ocr_page_indexes.append(analysis.page_index)
                continue

            target = word_dir / f"page_{page_number:05d}.docx"
            LOGGER.info("Page %d/%d: digital", page_number, total_pages)
            _emit(progress_callback, page_number, total_pages, "digital")

            try:
                _convert_digital_page(single_page_pdfs[analysis.page_index], target)
                _force_page_dimensions(target, analysis.width_mm, analysis.height_mm)
                page_docx_paths[analysis.page_index] = target
            except Exception as error:
                if not config.fallback_digital_to_ocr:
                    raise
                LOGGER.warning("Page %d: digital conversion failed; using OCR fallback (%s)", page_number, error)
                ocr_page_indexes.append(analysis.page_index)

        # Run the original high-quality PaddleOCR-VL route only for pages that
        # genuinely need it.
        if ocr_page_indexes:
            ocr_inputs = [single_page_pdfs[index] for index in ocr_page_indexes]
            LOGGER.info("OCR: processing %d page(s)", len(ocr_inputs))

            def ocr_progress(done: int, total: int, message: str) -> None:
                LOGGER.info("OCR: %d/%d result(s) ready", done, total)
                _emit(progress_callback, done, total, message)

            ocr_results = _best_ocr_results(
                ocr_inputs,
                config,
                quiet_library_logs=quiet_library_logs,
                progress_callback=ocr_progress,
            )

            for local_number, page_index in enumerate(ocr_page_indexes, start=1):
                analysis = analyses[page_index]
                page_number = page_index + 1
                input_page = single_page_pdfs[page_index].resolve()
                result = ocr_results.get(input_page)

                if result is None:
                    raise RuntimeError(f"No OCR result for page {page_number}.")

                target = word_dir / f"page_{page_number:05d}.docx"
                LOGGER.info("OCR export: page %d/%d", local_number, len(ocr_page_indexes))

                _save_ocr_result_to_docx(result, target, ocr_export_dir)
                _force_page_dimensions(target, analysis.width_mm, analysis.height_mm)
                page_docx_paths[page_index] = target

        ordered_page_docx = []
        for analysis in analyses:
            page_docx = page_docx_paths.get(analysis.page_index)
            if page_docx is None:
                raise RuntimeError(f"Page {analysis.page_index + 1} was not converted.")
            ordered_page_docx.append(page_docx)

        LOGGER.info("DOCX: combining %d page(s)", len(ordered_page_docx))
        _merge_docx_pages(ordered_page_docx, output_path, config)

    if config.validate_output:
        _validate_docx(output_path)

    elapsed = time.perf_counter() - started_at
    LOGGER.info("Done: %s (%.1f s)", output_path, elapsed)
    _emit(progress_callback, total_pages, total_pages, f"done in {elapsed:.1f}s")
    return output_path


def convert_many(
    input_pdfs: Iterable[str | os.PathLike],
    output_dir: str | os.PathLike | None = None,
    *,
    config: Optional[ConverterConfig] = None,
    continue_on_error: bool = True,
    quiet_library_logs: bool = True,
) -> list[BatchResult]:
    """Convert many PDFs inside ONE Python process.

    This is the preferred Python integration for a large queue because the OCR
    model is loaded at most once and then reused for every scan.
    """

    config = config or ConverterConfig()
    destination = Path(output_dir).expanduser().resolve() if output_dir else None
    if destination:
        destination.mkdir(parents=True, exist_ok=True)

    results: list[BatchResult] = []

    for input_value in input_pdfs:
        input_path = Path(input_value).expanduser().resolve()
        output_path = destination / (input_path.stem + ".docx") if destination else input_path.with_suffix(".docx")
        started = time.perf_counter()

        try:
            converted = convert_pdf_to_docx(
                input_path,
                output_path,
                config=config,
                quiet_library_logs=quiet_library_logs,
            )
            results.append(
                BatchResult(
                    input_pdf=str(input_path),
                    output_docx=str(converted),
                    success=True,
                    elapsed_seconds=time.perf_counter() - started,
                )
            )
        except Exception as error:
            LOGGER.error("Failed: %s | %s", input_path, error)
            results.append(
                BatchResult(
                    input_pdf=str(input_path),
                    output_docx=None,
                    success=False,
                    elapsed_seconds=time.perf_counter() - started,
                    error=str(error),
                )
            )
            if not continue_on_error:
                raise

    return results


# ---------------------------------------------------------------------------
# Persistent worker for integration from C#/C++/Java/Node/etc.
#
# Protocol: one JSON object per stdin line, one JSON response per stdout line.
# Our normal logs go to stderr, so stdout remains machine-readable.
# ---------------------------------------------------------------------------

def _worker_write(payload: dict) -> None:
    sys.stdout.write(json.dumps(payload, ensure_ascii=False) + "\n")
    sys.stdout.flush()


def run_worker(
    config: Optional[ConverterConfig] = None,
    *,
    preload_ocr: bool = False,
    quiet_library_logs: bool = True,
) -> int:
    """Run a long-lived conversion worker.

    Request examples:
        {"action":"ping"}
        {"action":"prepare_ocr"}
        {"action":"convert","input":"C:/a.pdf","output":"C:/a.docx"}
        {"action":"shutdown"}
    """

    config = config or ConverterConfig()

    if preload_ocr:
        try:
            prepare_ocr(config, quiet_library_logs=quiet_library_logs)
        except Exception as error:
            _worker_write({"ok": False, "event": "startup", "error": str(error)})
            return 1

    _worker_write({"ok": True, "event": "ready", "ocr_loaded": bool(_OCR_PIPELINE_CACHE)})

    for raw_line in sys.stdin:
        line = raw_line.strip()
        if not line:
            continue

        try:
            request = json.loads(line)
            if not isinstance(request, dict):
                raise ValueError("Worker request must be a JSON object.")

            action = str(request.get("action", "")).strip().lower()

            if action == "ping":
                _worker_write({"ok": True, "event": "pong", "ocr_loaded": bool(_OCR_PIPELINE_CACHE)})
                continue

            if action == "prepare_ocr":
                started = time.perf_counter()
                prepare_ocr(config, quiet_library_logs=quiet_library_logs)
                _worker_write({
                    "ok": True,
                    "event": "ocr_ready",
                    "elapsed_seconds": round(time.perf_counter() - started, 3),
                })
                continue

            if action == "convert":
                input_pdf = request.get("input")
                output_docx = request.get("output")
                if not input_pdf:
                    raise ValueError("convert request requires 'input'.")

                started = time.perf_counter()
                result = convert_pdf_to_docx(
                    input_pdf,
                    output_docx,
                    config=config,
                    quiet_library_logs=quiet_library_logs,
                )
                _worker_write({
                    "ok": True,
                    "event": "converted",
                    "input": str(Path(input_pdf)),
                    "output": str(result),
                    "elapsed_seconds": round(time.perf_counter() - started, 3),
                    "ocr_loaded": bool(_OCR_PIPELINE_CACHE),
                })
                continue

            if action == "shutdown":
                _worker_write({"ok": True, "event": "shutdown"})
                return 0

            raise ValueError(f"Unknown worker action: {action!r}")

        except Exception as error:
            _worker_write({"ok": False, "error": str(error)})

    return 0


# ---------------------------------------------------------------------------
# Command-line interface.
# ---------------------------------------------------------------------------

def _build_cli() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Hybrid PDF -> editable DOCX converter with persistent PaddleOCR-VL cache."
    )

    parser.add_argument("input_pdf", nargs="?")
    parser.add_argument("output_docx", nargs="?")

    parser.add_argument("--prepare-ocr", action="store_true", help="Download/load OCR models, then exit.")
    parser.add_argument("--worker", action="store_true", help="Run persistent JSON-lines worker on stdin/stdout.")
    parser.add_argument("--preload-ocr", action="store_true", help="With --worker, load OCR before reporting ready.")

    parser.add_argument("--mode", choices=["auto", "digital", "ocr"], default="auto")
    parser.add_argument("--device", default=None, help='Examples: "cpu", "gpu:0".')
    parser.add_argument("--engine", choices=["paddle", "transformers"], default=None)
    parser.add_argument("--cpu-threads", type=int, default=None)

    parser.add_argument("--no-unwarp", action="store_true")
    parser.add_argument("--no-orientation", action="store_true")
    parser.add_argument("--no-seal", action="store_true")
    parser.add_argument("--charts", action="store_true")
    parser.add_argument("--no-layout-retry", action="store_true")
    parser.add_argument(
        "--ocr-image-heavy-digital-pages",
        action="store_true",
        help="OCR digital pages merely because images occupy a large area (slower).",
    )

    parser.add_argument("--show-library-logs", action="store_true", help="Do not suppress Paddle/PaddleX console output.")
    parser.add_argument("--log-level", choices=["DEBUG", "INFO", "WARNING", "ERROR"], default="INFO")

    return parser


def _config_from_args(args) -> ConverterConfig:
    return ConverterConfig(
        mode=args.mode,
        device=args.device,
        engine=args.engine,
        cpu_threads=args.cpu_threads,
        use_doc_unwarping=not args.no_unwarp,
        use_doc_orientation_classify=not args.no_orientation,
        use_seal_recognition=not args.no_seal,
        use_chart_recognition=args.charts,
        retry_without_layout=not args.no_layout_retry,
        ocr_image_heavy_digital_pages=args.ocr_image_heavy_digital_pages,
    )


def main() -> int:
    args = _build_cli().parse_args()
    configure_logging(args.log_level)

    config = _config_from_args(args)
    quiet_library_logs = not args.show_library_logs

    try:
        if args.worker:
            return run_worker(
                config,
                preload_ocr=args.preload_ocr,
                quiet_library_logs=quiet_library_logs,
            )

        if args.prepare_ocr:
            started = time.perf_counter()
            LOGGER.info("Preparing OCR models. The first run may download several GB...")
            prepare_ocr(config, quiet_library_logs=quiet_library_logs)
            LOGGER.info("OCR preparation complete (%.1f s)", time.perf_counter() - started)
            return 0

        if not args.input_pdf:
            raise ValueError("input_pdf is required unless --prepare-ocr or --worker is used.")

        output = convert_pdf_to_docx(
            args.input_pdf,
            args.output_docx,
            config=config,
            quiet_library_logs=quiet_library_logs,
        )
        print(output)
        return 0

    except KeyboardInterrupt:
        LOGGER.warning("Interrupted by user")
        return 130
    except Exception as error:
        LOGGER.error("%s", error)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
