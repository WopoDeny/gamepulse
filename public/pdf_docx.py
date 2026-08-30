from __future__ import annotations

import tempfile
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

import fitz  # PyMuPDF
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Mm
from pdf2docx import Converter


APP_TITLE = "PDF to DOCX Converter"
TEXT_THRESHOLD = 20
SCAN_RENDER_DPI = 200


def is_real_pdf(file_path: str | Path) -> bool:
    """
    Check both the file extension and the PDF signature.
    """
    path = Path(file_path)

    if path.suffix.lower() != ".pdf":
        return False

    try:
        with path.open("rb") as file:
            return file.read(5) == b"%PDF-"
    except OSError:
        return False


def has_extractable_text(pdf_path: str | Path) -> bool:
    """
    Return True when the PDF contains a meaningful text layer.
    Scan-only PDFs usually return almost no text.
    """
    pdf = fitz.open(str(pdf_path))

    try:
        total_text = []

        for page in pdf:
            text = page.get_text("text").strip()

            if text:
                total_text.append(text)

            if sum(len(item) for item in total_text) >= TEXT_THRESHOLD:
                return True

        return False

    finally:
        pdf.close()


def unique_output_path(pdf_path: str | Path) -> Path:
    """
    Create a DOCX output path without silently overwriting an existing file.
    """
    pdf_path = Path(pdf_path)
    output = pdf_path.with_suffix(".docx")

    if not output.exists():
        return output

    counter = 1

    while True:
        candidate = pdf_path.with_name(f"{pdf_path.stem}_{counter}.docx")

        if not candidate.exists():
            return candidate

        counter += 1


def convert_text_pdf(pdf_path: str | Path, docx_path: str | Path) -> None:
    """
    Convert a regular PDF into an editable DOCX using pdf2docx.
    pdf2docx reads text blocks, font information, positions, tables,
    images and page layout from the PDF and reconstructs them in Word.
    """
    converter = Converter(str(pdf_path))

    try:
        converter.convert(str(docx_path))
    finally:
        converter.close()


def add_scan_page(
    document: Document,
    image_path: Path,
    width_mm: float,
    height_mm: float,
    first_page: bool,
) -> None:
    """
    Add one scan page as a full-page image while preserving page dimensions.
    """
    if first_page:
        section = document.sections[0]
    else:
        section = document.add_section(WD_SECTION.NEW_PAGE)

    section.page_width = Mm(width_mm)
    section.page_height = Mm(height_mm)

    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0

    run = paragraph.add_run()

    # A tiny safety reduction helps avoid Word pushing the image
    # to another page because of internal rounding.
    safe_width_mm = max(width_mm - 0.5, 1.0)

    run.add_picture(
        str(image_path),
        width=Mm(safe_width_mm),
    )


def convert_scan_pdf(pdf_path: str | Path, docx_path: str | Path) -> None:
    """
    Convert a scan-only PDF into DOCX by rendering each PDF page as an image.
    No OCR is performed, so the visual appearance is preserved.
    """
    pdf = fitz.open(str(pdf_path))
    document = Document()

    try:
        with tempfile.TemporaryDirectory(prefix="pdf_to_docx_") as temp_dir:
            temp_dir_path = Path(temp_dir)

            for page_index, page in enumerate(pdf):
                width_mm = page.rect.width * 25.4 / 72.0
                height_mm = page.rect.height * 25.4 / 72.0

                zoom = SCAN_RENDER_DPI / 72.0

                pixmap = page.get_pixmap(
                    matrix=fitz.Matrix(zoom, zoom),
                    alpha=False,
                )

                image_path = temp_dir_path / f"page_{page_index + 1}.png"
                pixmap.save(str(image_path))

                add_scan_page(
                    document=document,
                    image_path=image_path,
                    width_mm=width_mm,
                    height_mm=height_mm,
                    first_page=(page_index == 0),
                )

        document.save(str(docx_path))

    finally:
        pdf.close()


def convert_pdf(pdf_path: str | Path) -> Path:
    """
    Convert one PDF to DOCX.

    - Text-based PDF -> editable DOCX.
    - Scan-only PDF -> page images inside DOCX.
    """
    pdf_path = Path(pdf_path)

    if not pdf_path.exists():
        raise FileNotFoundError(f"File not found: {pdf_path}")

    if not is_real_pdf(pdf_path):
        raise ValueError(f"Not a valid PDF file: {pdf_path.name}")

    output_path = unique_output_path(pdf_path)

    if has_extractable_text(pdf_path):
        try:
            convert_text_pdf(pdf_path, output_path)
        except Exception:
            # Fallback: if editable reconstruction fails, preserve the document
            # visually instead of producing no output.
            if output_path.exists():
                output_path.unlink(missing_ok=True)

            convert_scan_pdf(pdf_path, output_path)
    else:
        convert_scan_pdf(pdf_path, output_path)

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("DOCX output was not created correctly.")

    return output_path


class PdfToDocxApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry("560x300")
        self.root.minsize(520, 280)

        self.status_var = tk.StringVar(value="Ready")

        self.build_ui()

    def build_ui(self) -> None:
        frame = ttk.Frame(self.root, padding=24)
        frame.pack(fill="both", expand=True)

        title = ttk.Label(
            frame,
            text="PDF to DOCX",
            font=("Segoe UI", 20, "bold"),
        )
        title.pack(pady=(5, 8))

        subtitle = ttk.Label(
            frame,
            text=(
                "Text PDFs are converted to editable Word documents.\n"
                "Scan-only PDFs are preserved as page images."
            ),
            justify="center",
        )
        subtitle.pack(pady=(0, 20))

        self.convert_button = ttk.Button(
            frame,
            text="Select PDF files",
            command=self.select_and_convert,
        )
        self.convert_button.pack(ipadx=18, ipady=8)

        self.progress = ttk.Progressbar(
            frame,
            mode="determinate",
            length=360,
        )
        self.progress.pack(pady=(22, 8))

        status = ttk.Label(
            frame,
            textvariable=self.status_var,
            anchor="center",
        )
        status.pack(fill="x")

    def select_and_convert(self) -> None:
        selected_files = filedialog.askopenfilenames(
            title="Select PDF files",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("All files", "*.*"),
            ],
        )

        if not selected_files:
            return

        self.convert_button.config(state="disabled")
        self.progress["maximum"] = len(selected_files)
        self.progress["value"] = 0

        created_files = []
        errors = []

        try:
            for index, file_path in enumerate(selected_files, start=1):
                source = Path(file_path)
                self.status_var.set(f"Converting: {source.name}")
                self.root.update_idletasks()

                try:
                    output = convert_pdf(source)
                    created_files.append(output)
                except Exception as error:
                    errors.append(f"{source.name}: {error}")

                self.progress["value"] = index
                self.root.update_idletasks()

        finally:
            self.convert_button.config(state="normal")

        if created_files:
            self.status_var.set(
                f"Done. Created {len(created_files)} DOCX file(s)."
            )
        else:
            self.status_var.set("Conversion failed.")

        if created_files and not errors:
            result_text = "\n".join(str(path) for path in created_files)

            messagebox.showinfo(
                "Conversion complete",
                f"Created:\n\n{result_text}",
            )

        elif created_files and errors:
            result_text = "\n".join(str(path) for path in created_files)
            error_text = "\n".join(errors)

            messagebox.showwarning(
                "Conversion completed with warnings",
                f"Created:\n\n{result_text}\n\nErrors:\n{error_text}",
            )

        else:
            messagebox.showerror(
                "Conversion failed",
                "\n".join(errors) or "Unknown error.",
            )


def main() -> None:
    root = tk.Tk()
    PdfToDocxApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
